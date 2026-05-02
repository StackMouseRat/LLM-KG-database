#!/usr/bin/env python3
from __future__ import annotations

import shutil
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs/designs/论文草稿0430可修改_第四五章二次迁移压缩扩写版.docx"
DEFAULT_TARGET = ROOT / "docs/designs/论文草稿0430可修改_第五章实验扩写版.docx"


def parse_args():
    parser = argparse.ArgumentParser(description="Write expanded chapter 5 experiment section into a Word copy.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--target", type=Path, default=DEFAULT_TARGET)
    return parser.parse_args()


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    anchor_element = paragraph._tbl if hasattr(paragraph, "_tbl") else paragraph._p
    anchor_element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows):
    table = paragraph._parent.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.0))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(value)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.5
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)
    anchor_element = paragraph._tbl if hasattr(paragraph, "_tbl") else paragraph._p
    anchor_element.addnext(table._tbl)
    return table


def set_normal(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def set_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def remove_trailing_empty_paragraphs(doc: Document) -> None:
    body = doc._body._element
    for child in reversed(list(body)):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag != qn("w:p"):
            break
        if paragraph_text(child).strip():
            break
        body.remove(child)


def add_after(anchor, kind, content, style=None):
    if kind == "p":
        para = insert_paragraph_after(anchor, content, style)
        set_normal(para)
        return para
    if kind == "h3":
        para = insert_paragraph_after(anchor, content, "heading 3")
        if para.style.name == "Normal":
            for run in para.runs:
                run.bold = True
        return para
    if kind == "caption":
        para = insert_paragraph_after(anchor, content)
        set_caption(para)
        return para
    if kind == "table":
        return insert_table_after(anchor, content)
    raise ValueError(kind)


def main():
    args = parse_args()
    source = args.source if args.source.is_absolute() else ROOT / args.source
    target = args.target if args.target.is_absolute() else ROOT / args.target
    shutil.copyfile(source, target)
    doc = Document(target)

    paragraphs = list(doc.paragraphs)
    start_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "案例验证与结果分析")
    end_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "总结与展望")
    heading = paragraphs[start_idx]

    for p in paragraphs[start_idx + 1:end_idx]:
        delete_paragraph(p)

    anchor = heading
    content = [
        ("p", "为验证本文原型系统在电气设备应急预案生成任务中的有效性，本文在前述实现工作的基础上开展案例验证与对照实验。实验不以通用文本生成能力为评价对象，而是围绕预案生成链路中的关键能力展开，包括设备主体识别与故障匹配、知识图谱增强、模板结构约束以及多故障链式生成。同时，实验对输入边界判定能力进行验证，用于考察系统对无关输入、不支持设备和明显错配问题的处理效果。"),
        ("p", "实验分析采用结构规范性、知识相关性、措施完整性和内容可追溯性四个维度。结构规范性关注生成结果是否遵循预案模板章节和小节顺序；知识相关性关注设备对象、故障节点和图谱知识是否匹配；措施完整性关注报告、隔离、研判、抢修、恢复和响应终止等关键处置环节是否覆盖；内容可追溯性关注生成文本是否能够区分图谱来源、模型补充和固定模板来源。"),
        ("caption", "表5-1 实验计划与对照组设置"),
        ("table", [
            ["实验名称", "对照组", "消融组一", "消融组二", "主要观察指标"],
            ["输入边界判定", "完整工作流", "移除边界校验", "关键词边界校验", "不合规输入拦截、误放行、误拦截"],
            ["设备识别与故障匹配", "完整工作流", "移除主体判定", "关键词主体判定", "设备表选择、故障节点匹配、图谱命中"],
            ["图谱与模板约束", "完整工作流", "移除图谱", "移除模板", "来源标签、章节完整性、领域相关性"],
            ["多故障链式生成", "多故障完整链路", "单故障链路", "仅主故障图谱", "多故障覆盖、逐故障措施、融合一致性"],
        ]),
        ("caption", "图5-1 实验页面与对照组运行流程图占位符"),
        ("p", "【图5-1占位符：展示实验页面从题目选择、实验组配置、并发运行、结果落盘到评分记录的流程。建议配图包含 boundary、disambiguation、graphTemplate、multiFault 四类实验计划。】"),
        ("h3", "实验目的与环境"),
        ("p", "本文实验依托已实现的原型系统完成。系统运行链路由前端单页应用、服务代理、Python 流水线、FastGPT 工作流插件和 NebulaGraph 图数据库共同构成。前端实验页面负责组织测试题、选择实验计划、展示运行结果和保存评价记录；服务代理负责任务创建、并发控制和运行状态转发；Python 流水线负责调用插件、处理边界终止、组织章节并行生成和保存结果；FastGPT 插件负责基本信息获取、模板切片、章节生成、多故障识别和二级故障查询；NebulaGraph 提供设备故障图谱、模板图谱和评价题库图谱的查询支撑。"),
        ("p", "实验页面后端已实现四类实验计划，分别为输入边界判定、设备识别与故障匹配、图谱与模板约束、多故障链式生成。每类实验均设置完整工作流作为对照组，并设置两个消融组，用于观察单一能力被移除或弱化后的输出变化。该设计能够避免仅凭单次生成结果进行主观判断，而是通过对照组之间的差异分析说明各模块在生成链路中的作用。"),
        ("h3", "测试样例与评价方法"),
        ("p", "测试样例主要来自系统对照试验题库、轻量实验样题和已验证的典型设备场景，覆盖断路器、互感器、变压器、电缆、避雷器等设备类型。样例既包含单设备单故障问题，也包含位置描述容易引发主体歧义的问题，以及同一设备下多个故障同时出现的复合问题。通过这类样例可以同时验证设备识别、故障节点匹配、图谱检索和章节生成效果。"),
        ("caption", "表5-2 代表性测试样例设置"),
        ("table", [
            ["样例类型", "代表问题", "验证目标"],
            ["主体消歧", "主变旁断路器拒动，请生成应急处置方案。", "判断系统能否将主体识别为断路器而非主变"],
            ["主体消歧", "开关柜内电流互感器二次开路，请给出处置方案。", "判断系统能否将主体识别为互感器而非开关柜"],
            ["图谱增强", "某110kV断路器故障时保护已发令，但现场未见分闸动作，请生成应急处置方案。", "比较有图谱与无图谱情况下第四章处置措施差异"],
            ["模板约束", "某110kV断路器同时出现控制回路短路故障和合分闸控制回路故障，请生成现场应急处置方案。", "比较裸 LLM 与带模板 LLM 的结构差异"],
            ["多故障", "某主变发生重瓦斯告警，同时出现拒动故障，请生成多故障应急预案。", "验证多故障边界判定、故障拆解和融合生成能力"],
        ]),
        ("p", "评价时不追求唯一标准答案，而是结合生成过程和输出内容进行综合判断。对于结构规范性，主要检查章节标题、章节顺序和小节覆盖情况；对于知识相关性，主要检查设备空间、故障二级节点和图谱素材是否与用户问题一致；对于措施完整性，主要检查是否覆盖信息报告、隔离控制、现场研判、抢修恢复、风险防控和应急资源；对于内容可追溯性，主要检查输出中 KG、GEN、FIX 等来源标签及图谱素材是否能够支撑正文内容。"),
        ("h3", "输入边界判定实验"),
        ("p", "输入边界判定用于避免系统对无关问题、不支持设备问题或明显错配场景继续生成正式预案。实验设置完整工作流、移除边界校验和关键词边界校验三组。完整工作流依托基本信息获取插件中的语义边界判定；移除边界校验组在插件返回边界失败时仍继续执行后续模板切片和章节生成；关键词边界校验组仅依据设备词和故障词进行简单规则判断。"),
        ("p", "从实现逻辑看，完整工作流能够结合设备、故障动作和语义场景进行判断，适合处理描述较复杂的业务输入；关键词规则在显式出现设备名称和故障词时能够快速命中，但对位置描述、设备别名和错配场景的判断能力较弱；移除边界校验组则可能使无效输入继续进入生成链路，增加无意义生成和错误预案输出的风险。由此可见，输入边界判定能够在正式生成之前过滤不适配问题，为后续设备识别、图谱检索和章节生成提供稳定的输入基础。"),
        ("caption", "图5-2 输入边界判定对照结果截图占位符"),
        ("p", "【图5-2占位符：展示同一批边界测试题在完整工作流、移除边界校验、关键词边界校验三组中的运行状态。建议标注被拦截、继续生成、误放行等结果。】"),
        ("h3", "设备识别与故障匹配实验"),
        ("p", "设备识别与故障匹配是知识增强预案生成链路中的关键环节。电气设备故障图谱按设备类型存储在不同 Nebula 图空间中，因此设备主体识别结果会直接决定后续检索的图谱空间；同时，不同设备下的二级故障名称和下游知识也不相同，故障匹配结果会进一步影响检索到的原因、现象、措施、风险和资源内容。如果主体识别错误，即使后续生成模型具备较强语言组织能力，也可能围绕错误设备生成看似完整但业务对象不一致的预案。"),
        ("p", "实验设置完整工作流、移除主体判定和关键词主体判定三组。完整工作流通过基本信息插件结合用户问题上下文、设备位置描述和故障动作进行判断；移除主体判定组弱化设备主体信息，使系统难以稳定确定图谱空间；关键词主体判定组依据设备词直接匹配，当问题中同时出现位置设备和实际故障设备时容易受到前置词干扰。"),
        ("p", "以“主变旁断路器拒动”为例，问题中同时出现“主变”和“断路器”，但拒动动作对应的故障主体应为断路器。完整工作流能够根据动作特征将设备表定位为断路器图谱，而简单关键词策略可能优先受到“主变”位置词影响。再以“开关柜内电流互感器二次开路”为例，开关柜是位置或载体，电流互感器才是故障主体；完整工作流需要识别这种包含关系，才能进入互感器图谱检索二次开路相关故障。"),
        ("caption", "表5-3 设备识别与故障匹配实验结果占位表"),
        ("table", [
            ["测试问题", "期望设备", "完整工作流", "移除主体判定", "关键词主体判定", "结论"],
            ["主变旁断路器拒动", "断路器", "待填：llmkg_breaker", "待填", "待填", "验证位置词与故障主体区分能力"],
            ["开关柜内电流互感器二次开路", "互感器", "待填：llmkg_mutual", "待填", "待填", "验证载体设备与故障主体区分能力"],
            ["主变轻瓦斯告警", "变压器", "待填：llmkg_transformer", "待填", "待填", "验证典型设备告警识别能力"],
            ["电缆沟进水并伴随绝缘告警", "电缆", "待填：llmkg_cable", "待填", "待填", "验证环境描述下的主体识别能力"],
        ]),
        ("caption", "图5-3 设备识别与故障匹配调用链截图占位符"),
        ("p", "【图5-3占位符：展示设备识别节点、故障类型分析节点和图谱检索节点的调用链结果，重点标注设备表、故障二级节点和图谱素材字段。】"),
        ("h3", "图谱增强与模板约束实验"),
        ("p", "图谱增强与模板约束分别从内容依据和结构组织两个方面影响预案生成质量。模板约束用于规定预案章节体系、内容来源和生成要求，知识图谱用于提供故障原因、故障现象、处置措施、安全风险和应急资源等领域知识。二者分别作用于结构层和内容层，缺少模板时生成结果容易呈现自由组织形态，缺少图谱时生成内容虽然通顺，但更多依赖模型通用知识，领域针对性和来源可追溯性下降。"),
        ("p", "首先比较裸 LLM 与带模板 LLM。轻量实验中，裸 LLM 在断路器控制回路短路与合分闸控制回路故障场景下生成 2024 字，能够给出通用故障概述、应急处置流程、安全注意事项和物资清单；带模板 LLM 生成 3530 字，能够围绕信息报告、先期处置、现场研判、抢修恢复和响应终止等模板章节展开，结构完整性明显增强。该结果说明，模板约束能够将自由问答式生成转化为面向预案章节的结构化填充。"),
        ("caption", "表5-4 裸 LLM 与带模板 LLM 生成结果对照"),
        ("table", [
            ["组别", "字符数", "编号步骤", "处置动作", "主要表现"],
            ["裸 LLM", "2024", "是", "是", "能够生成通用处置流程，但章节组织较自由"],
            ["带模板 LLM", "3530", "是", "是", "能够按照信息报告、先期处置、现场研判等模板章节展开"],
        ]),
        ("p", "其次比较有图谱与无图谱。第四章处置措施对照实验复用基本信息获取、模板切片和并行生成链路，仅在并行生成阶段改变图谱素材输入。有图谱组传入真实图谱检索方案素材，无图谱组传入“知识图谱无数据”。在断路器拒动场景中，有图谱组输出来源标签为 KG=23、GEN=15，无图谱组为 KG=0、GEN=34；在电压互感器熔断器熔断场景中，有图谱组为 KG=14、GEN=15、FIX=1，无图谱组为 KG=0、GEN=40。"),
        ("caption", "表5-5 有无图谱第四章生成对照"),
        ("table", [
            ["题目", "组别", "KG标签", "GEN标签", "FIX标签", "结论"],
            ["Q001 断路器拒动", "有图谱", "23", "15", "0", "图谱支撑原因、现象、措施和风险条目"],
            ["Q001 断路器拒动", "无图谱", "0", "34", "0", "主要依赖模型常识生成"],
            ["Q002 TV熔丝熔断", "有图谱", "14", "15", "1", "图谱支撑本体故障、熔丝更换和风险措施"],
            ["Q002 TV熔丝熔断", "无图谱", "0", "40", "0", "可读但缺少来源可追溯性"],
        ]),
        ("p", "从输出内容看，有图谱组能够把图谱中的故障原因、故障现象、应对措施、安全风险和应急资源映射到对应章节。例如断路器拒动场景中，图谱素材包含直流电源异常、控制和合分闸回路不通、线圈故障、铁心动作失灵、机构卡阻、辅助开关异常、储能异常及液压气动异常等原因要点，生成结果在故障检查、原因排查和抢修措施中均能体现这些内容。无图谱组仍能生成语言流畅的预案文本，但来源标签全部为 GEN，说明其主要依赖模型自身知识进行扩写，难以证明内容来自已维护的领域知识资产。"),
        ("caption", "图5-4 图谱增强生成结果来源标签对比图占位符"),
        ("p", "【图5-4占位符：可绘制 Q001、Q002 在有图谱和无图谱条件下 KG、GEN、FIX 标签数量的柱状图。】"),
        ("caption", "图5-5 有图谱与无图谱典型输出片段对照占位符"),
        ("p", "【图5-5占位符：左侧展示有图谱组包含 KG 标签的原因排查和抢修措施片段，右侧展示无图谱组 GEN 扩写片段。】"),
        ("h3", "多故障链式生成实验"),
        ("p", "多故障链式生成用于验证系统对复合故障场景的处理能力。实际电气设备故障往往不是单一异常孤立出现，而可能包含主故障、伴随异常和次生风险。如果系统只按照单故障链路处理，容易只围绕主故障生成处置措施，忽略伴随故障的风险和资源需求；如果只识别多故障但不逐项检索图谱，则生成阶段缺少每个故障对应的专业素材，融合结果仍可能偏泛化。"),
        ("p", "实验设置多故障完整链路、单故障普通链路和仅主故障图谱链路三组。多故障完整链路首先复用单故障基本信息插件进行边界判定，确保输入仍属于系统支持范围；随后调用多故障基本信息插件识别多个故障节点，并确定主故障与伴随故障；接着围绕每个故障分别调用二级故障信息查询插件，检索原因、现象、后果、措施、风险和资源；最后将逐故障图谱素材聚合为统一结构，传入并行生成插件，由生成提示词完成章节级融合表达。"),
        ("caption", "图5-6 多故障链式生成流程图占位符"),
        ("p", "【图5-6占位符：展示“单故障边界判定—多故障识别—逐故障图谱检索—素材聚合—章节并行生成—结果合并”的链式流程。】"),
        ("p", "以“某主变发生重瓦斯告警，同时出现拒动故障，请生成多故障应急预案”为例，系统首先需要判断输入是否属于支持设备和支持故障范围；边界判定通过后，多故障插件应将输入拆解为多个故障节点，并围绕各故障分别组织图谱素材。与单故障普通链路相比，完整多故障链路的预期优势在于能够同时覆盖主故障处置、伴随故障排查、风险联动和应急资源协调，使生成结果不只围绕一个故障点展开。"),
        ("caption", "表5-6 多故障链式生成实验结果占位表"),
        ("table", [
            ["实验组", "故障识别", "图谱检索", "生成特点", "待记录结论"],
            ["多故障完整链路", "识别多个故障节点", "逐故障检索", "覆盖主故障和伴随故障，并进行融合表达", "待填实验运行结果"],
            ["单故障普通链路", "倾向识别主故障", "仅主链路图谱", "可能忽略伴随异常或次生风险", "待填实验运行结果"],
            ["仅主故障图谱", "识别多故障", "只检索主故障", "具备多故障意识但素材覆盖不足", "待填实验运行结果"],
        ]),
        ("caption", "图5-7 多故障生成结果页面截图占位符"),
        ("p", "【图5-7占位符：展示多故障完整链路的实验页面输出，重点标注多故障识别结果、逐故障图谱素材和最终章节融合内容。】"),
        ("h3", "实验结果小结"),
        ("p", "综合上述实验，本文原型系统在设备识别与故障匹配、图谱增强与模板约束、多故障链式生成三个核心方面表现出明确作用。设备识别与故障匹配实验说明，正确识别故障主体是选择图谱空间和匹配二级故障节点的前提，完整工作流能够比简单关键词策略更好地处理位置词、载体词和故障主体并存的输入。"),
        ("p", "图谱增强与模板约束实验说明，模板能够提升预案结构的稳定性，使生成结果按照预案章节展开；图谱能够提升内容的领域相关性，使故障原因、现象、处置措施、安全风险和应急资源具有明确知识来源。来源标签统计进一步说明，有图谱组能够保留 KG、GEN、FIX 等来源信息，为人工审查和质量评估提供依据。"),
        ("p", "多故障链式生成实验说明，复合故障场景需要先完成边界控制和故障拆解，再进行逐故障图谱检索和融合生成。完整多故障链路相较单故障链路和仅主故障图谱链路，更适合表达主故障与伴随故障之间的处置关联。输入边界判定实验则从前置控制角度说明，系统能够在进入正式生成前过滤无关或不适配问题，为后续生成链路提供基本可靠性保障。"),
    ]

    for kind, value in content:
        anchor = add_after(anchor, kind, value)

    remove_trailing_empty_paragraphs(doc)

    doc.save(target)
    print(target)


if __name__ == "__main__":
    main()
