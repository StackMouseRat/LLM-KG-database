#!/usr/bin/env python3
from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs/designs/论文草稿0502可修改_第五章实验扩写版.docx"
DEFAULT_TARGET = ROOT / "docs/designs/论文草稿0502可修改_第六章总结展望版.docx"


def parse_args():
    parser = argparse.ArgumentParser(description="Rewrite chapter 6 summary and outlook.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--target", type=Path, default=DEFAULT_TARGET)
    return parser.parse_args()


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def set_normal(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def add_after(anchor, kind, text):
    if kind == "h2":
        para = insert_paragraph_after(anchor, text, "heading 2")
        return para
    if kind == "p":
        para = insert_paragraph_after(anchor, text)
        set_normal(para)
        return para
    raise ValueError(kind)


def main():
    args = parse_args()
    source = args.source if args.source.is_absolute() else ROOT / args.source
    target = args.target if args.target.is_absolute() else ROOT / args.target
    shutil.copyfile(source, target)
    doc = Document(target)

    paragraphs = list(doc.paragraphs)
    start_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "总结与展望")
    end_idx = next(i for i, p in enumerate(paragraphs[start_idx + 1:], start_idx + 1) if p.text.strip() == "参考文献")
    heading = paragraphs[start_idx]

    for p in paragraphs[start_idx + 1:end_idx]:
        delete_paragraph(p)

    content = [
        ("h2", "全文总结"),
        ("p", "本文面向电气设备应急预案编制中人工参与环节较多、知识组织较为分散、故障场景差异较大和生成结果需要可追溯支撑等特点，研究并实现了一种基于大语言模型与知识图谱的应急预案智能生成方法。全文围绕预案模板体系、设备故障知识图谱、知识增强生成流程和原型系统实现展开，形成了从知识建模、图谱入库、工作流编排到前端验证的完整技术闭环。"),
        ("p", "在预案结构建模方面，本文首先分析电气设备应急预案的章节规律和内容组成，将预案文本拆分为相对稳定的模板结构与面向具体故障场景动态生成的内容单元。通过对章节标题、内容来源、固定文本和生成要求进行结构化表达，本文将传统整篇式预案编写转化为面向章节槽位的生成任务。与仅依赖大语言模型直接生成整篇文本的方法相比，该方式能够在生成前明确章节边界和内容约束，使输出更符合应急预案的规范化组织要求。"),
        ("p", "在知识组织方面，本文围绕断路器、电缆、变压器、互感器等典型设备构建设备故障知识图谱，采用“设备入口—故障分类—具体故障—下游知识”的层级结构组织故障原因、故障现象、故障后果、处置措施、安全风险和应急资源等知识要素。与已有研究中侧重事故文本检索、单一设备知识库或通用应急知识组织的方法相比，本文将预案生成所需的专业知识直接映射到模板内容来源和章节生成字段，使知识图谱不只是检索背景材料，而是参与预案正文生成和来源追踪的结构化依据。"),
        ("p", "在生成流程设计方面，本文构建了由基本信息获取、模板切片、并行章节生成、多故障识别和二级故障查询等模块组成的知识增强工作流。系统首先识别用户输入中的设备主体、故障节点和场景要素，再根据故障类型检索图谱素材，并在模板约束下分章节生成预案正文。针对复合故障场景，本文进一步设计了多故障链式生成机制，通过故障拆解、逐故障图谱检索和统一素材聚合，使生成过程能够覆盖主故障、伴随异常和风险联动关系。与单轮问答式生成或简单 RAG 拼接方式相比，该流程在任务分解、过程控制和复杂场景组织方面具有更强的可控性。"),
        ("p", "在系统实现方面，本文完成了包含前端单页应用、服务代理、Python 流水线、FastGPT 工作流和 NebulaGraph 图数据库的原型系统。系统支持预案生成、图谱溯源、模板查看、质量评估和对比实验等功能，并通过服务接入层统一组织认证、生成请求、流式返回、实验运行和结果复用。该实现将大语言模型能力、知识图谱查询和工程化运行控制结合起来，使预案生成过程不再停留在离线提示词实验层面，而具备了可运行、可观测、可复盘的系统形态。"),
        ("p", "在实验验证方面，本文围绕设备识别与故障匹配、图谱增强与模板约束、多故障链式生成等关键能力开展案例验证和对照实验。实验结果表明，模板约束能够提升生成结果的结构规范性，知识图谱能够增强内容的领域相关性和可追溯性，多故障链式处理能够改善复合场景下的故障覆盖和融合表达。总体来看，本文工作的创新点主要体现在三个方面：一是将预案模板、图谱知识和大语言模型生成过程进行细粒度耦合；二是面向多设备、多故障场景设计了可扩展的图谱检索与链式生成机制；三是构建了从知识建设到实验验证的一体化原型系统，为电气设备应急预案智能生成提供了可实现的技术路径。"),
        ("h2", "展望"),
        ("p", "后续研究可在现有原型基础上继续深化知识资产建设。随着设备类型和故障样本的增加，设备故障图谱可以进一步扩展到更多电压等级、更多设备型号和更细粒度的故障场景，并在运维规程、历史缺陷、检修记录和现场处置报告之间建立更加稳定的知识映射关系。通过持续补充和校验图谱节点、关系及模板槽位，系统能够获得更充分的领域知识支撑，从而适应更多实际生产场景。"),
        ("p", "在故障识别和场景理解方面，后续可继续优化复杂语句下的主体消歧、故障层级匹配和跨故障关系判断。电力现场描述常常包含位置设备、故障设备、保护动作和运行方式等多类信息，部分输入还会同时涉及主故障、伴随告警和次生风险。未来可结合更丰富的提示策略、规则校验和图谱约束机制，使系统在复杂输入下保持较稳定的设备识别和故障匹配效果，并进一步提升多故障素材组织的准确性。"),
        ("p", "在生成质量评估方面，后续可完善更系统的量化评价体系。本文已经从结构规范性、知识相关性、措施完整性和内容可追溯性等维度开展验证，后续可结合更多测试题、专家评分和自动评估结果，形成覆盖不同设备、不同故障类型和不同生成模式的评测集。通过更充分的对照实验，可以进一步分析模板、图谱、案例知识和多故障链路对生成质量的影响，为后续模型选择和流程优化提供依据。"),
        ("p", "在工程应用方面，后续可增强原型系统与实际业务流程的衔接能力。例如，可进一步完善前端图谱交互、模板版本管理、运行记录检索和多人协同审查等功能，使生成结果能够更方便地进入人工审核、预案修订和知识回流流程。同时，也可以探索与企业现有运维系统、缺陷管理系统和应急指挥平台的接口集成，使本文提出的方法从研究型原型逐步向更稳定的业务支撑工具演进。"),
        ("p", "总体而言，基于大语言模型与知识图谱的电气设备应急预案生成仍具有较大的拓展空间。随着领域知识持续积累、模型能力不断提升以及工程接口逐步完善，该类方法有望在预案辅助编制、应急知识复用、故障处置建议生成和运维经验沉淀等方面发挥更大作用，为电力设备应急管理的数字化和智能化提供持续支撑。"),
    ]

    anchor = heading
    for kind, text in content:
        anchor = add_after(anchor, kind, text)

    doc.save(target)
    print(target)


if __name__ == "__main__":
    main()
